import os
import re
import hashlib
import requests
from dotenv import load_dotenv
from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse, StreamingResponse
from .file_manager import (
    create_file_record, save_uploaded_file, add_file_record,
    get_all_files, get_file_record, delete_file_record,
    delete_physical_file
)

# 加载配置
load_dotenv()

router = APIRouter()

# ---------------------- 全局懒加载变量 ----------------------
_vector_model = None
_chroma_client = None
_collection = None


def _get_model():
    """懒加载：首次调用时初始化向量模型"""
    global _vector_model
    if _vector_model is None:
        os.environ["TRANSFORMERS_OFFLINE"] = "1"
        os.environ["HF_HUB_OFFLINE"] = "1"
        from sentence_transformers import SentenceTransformer
        snapshot_path = r"C:\Users\Administrator\.cache\huggingface\hub\models--BAAI--bge-large-zh-v1.5\snapshots\79e7739b6ab944e86d6171e44d24c997fc1e0116"
        if os.path.exists(snapshot_path):
            print(f"[INFO] 从快照加载：{snapshot_path}")
            _vector_model = SentenceTransformer(snapshot_path)
            print(f"[INFO] OK 模型加载成功")
        else:
            model_name = os.getenv("VECTOR_MODEL", "BAAI/bge-large-zh-v1.5")
            print(f"[INFO] 使用模型名称加载：{model_name}")
            _vector_model = SentenceTransformer(model_name)
            print(f"[INFO] OK 模型加载成功")
    return _vector_model


def _get_collection():
    """懒加载：首次调用时初始化向量库"""
    global _chroma_client, _collection
    if _collection is None:
        import chromadb
        _chroma_client = chromadb.PersistentClient(path=os.getenv("CHROMA_DB_PATH", "./dba_knowledge_db"))
        _collection = _chroma_client.get_or_create_collection(name="dba_knowledge")
    return _collection


# ---------------------- 1. DBA 文档脱敏函数 ----------------------
def desensitize_dba_text(text: str) -> str:
    """DBA 文档专属脱敏：IP、账号、敏感表名、密码"""
    text = re.sub(r'\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}', 
                  lambda x: re.sub(r'\d{1,3}', '×', x.group(0), count=2), text)
    text = re.sub(r'(root|admin|dba)_?\w*:\w+', r'****:****', text)
    text = re.sub(r'(user|order|payment|account)_\w+', r'tb_××××', text)
    text = re.sub(r':\d{4,5}', r':××××', text)
    return text


# ---------------------- 2. 文档解析函数 ----------------------
def parse_document(file: UploadFile) -> str:
    """解析上传的文件（txt/md/sql/pdf/docx）"""
    try:
        from pypdf import PdfReader
        import docx
        ext = file.filename.split(".")[-1].lower()
        content = ""
        if ext in ["txt", "md", "sql"]:
            content = file.file.read().decode("utf-8")
        elif ext == "pdf":
            reader = PdfReader(file.file)
            content = "\n".join([page.extract_text() for page in reader.pages])
        elif ext == "docx":
            doc = docx.Document(file.file)
            content = "\n".join([para.text for para in doc.paragraphs])
        else:
            raise HTTPException(status_code=400, detail=f"不支持的文件格式：{ext}")
        return content
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件解析失败：{str(e)}")


def parse_document_from_content(content: bytes, filename: str) -> str:
    """从文件内容解析文本（txt/md/sql/pdf/docx）"""
    try:
        from io import BytesIO
        from pypdf import PdfReader
        import docx
        ext = filename.split(".")[-1].lower()
        text = ""
        if ext in ["txt", "md", "sql"]:
            text = content.decode("utf-8")
        elif ext == "pdf":
            reader = PdfReader(BytesIO(content))
            text = "\n".join([page.extract_text() for page in reader.pages])
        elif ext == "docx":
            doc = docx.Document(BytesIO(content))
            text = "\n".join([para.text for para in doc.paragraphs])
        else:
            raise HTTPException(status_code=400, detail=f"不支持的文件格式：{ext}")
        return text
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件解析失败：{str(e)}")


# ---------------------- 3. 调用豆包 API 解析文档 ----------------------
def call_doubao_api(text: str) -> str:
    """调用豆包 API 解析 DBA 文档，提取结构化知识点"""
    api_key = os.getenv("DOUBAO_API_KEY")
    endpoint = os.getenv("DOUBAO_API_ENDPOINT")
    if not api_key or not endpoint:
        raise HTTPException(status_code=500, detail="豆包 API 配置未完成")
    prompt = f"""你是资深 DBA，请解析以下文档内容，完成 2 件事：
1. 提取核心知识点（故障现象、原因、排查步骤、解决命令、注意事项）；
2. 用简洁的结构化格式输出（分点说明，关键命令用代码块标注）。

文档内容：
{text}
"""
    try:
        response = requests.post(
            endpoint,
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            json={
                "model": "doubao-1-5-pro-32k-250115",
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.1
            }
        )
        response.raise_for_status()
        return response.json()["choices"][0]["message"]["content"]
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"豆包 API 调用失败：{str(e)}")


# ---------------------- 4. 向量库增量入库 ----------------------
def add_to_chroma(text: str, file_name: str, chunk_prefix: str = ""):
    """解析结果存入 Chroma 向量库（增量入库，自动去重）"""
    vector_model = _get_model()
    collection = _get_collection()
    
    lines = text.split('\n')
    chunks = []
    current_chunk = []
    current_length = 0
    max_chunk_size = 300

    for line in lines:
        line = line.strip()
        if not line:
            continue
        if any(keyword in line.upper() for keyword in ['SELECT', 'INSERT', 'UPDATE', 'DELETE', 'FROM', 'WHERE']):
            if current_chunk:
                chunks.append('\n'.join(current_chunk))
                current_chunk = []
                current_length = 0
            chunks.append(line)
        else:
            if current_length + len(line) > max_chunk_size and current_chunk:
                chunks.append('\n'.join(current_chunk))
                current_chunk = [line]
                current_length = len(line)
            else:
                current_chunk.append(line)
                current_length += len(line)
    if current_chunk:
        chunks.append('\n'.join(current_chunk))

    chunk_count = 0
    for chunk in chunks:
        if not chunk.strip():
            continue
        chunk_id = hashlib.md5(f"{chunk_prefix}{chunk}".encode()).hexdigest()
        if collection.get(ids=[chunk_id])["ids"]:
            continue
        embedding = vector_model.encode(chunk).tolist()
        collection.add(
            ids=[chunk_id],
            embeddings=[embedding],
            documents=[chunk],
            metadatas=[{
                "file_name": file_name, 
                "type": "dba_knowledge",
                "chunk_size": len(chunk),
                "has_sql": any(k in chunk.upper() for k in ['SELECT', 'INSERT', 'UPDATE', 'DELETE']),
                "source_type": "raw" if chunk_prefix == "raw_" else "ai_parsed"
            }]
        )
        chunk_count += 1
    return chunk_count


# ---------------------- 5. 上传文件 + 搭建知识库 ----------------------
@router.post("/upload-and-build")
async def upload_and_build(file: UploadFile = File(...)):
    print(f"\n{'='*60}")
    print(f"[UPLOAD] 收到上传请求")
    print(f"[UPLOAD] 文件名：{file.filename}")
    print(f"{'='*60}")
    try:
        file_content = await file.read()
        print(f"[UPLOAD] 步骤 1/5: 解析文件...")
        raw_content = parse_document_from_content(file_content, file.filename)
        print(f"[UPLOAD] OK 文件解析成功，内容长度：{len(raw_content)}")
        print(f"[UPLOAD] 步骤 2/5: 脱敏处理...")
        desensitized_content = desensitize_dba_text(raw_content)
        print(f"[UPLOAD] OK 脱敏完成")
        print(f"[UPLOAD] 步骤 3/5: 原始内容向量入库...")
        raw_chunk_count = add_to_chroma(desensitized_content, file.filename, chunk_prefix="raw_")
        print(f"[UPLOAD] OK 原始内容入库，片段数：{raw_chunk_count}")
        print(f"[UPLOAD] 步骤 4/5: 调用豆包 AI 解析...")
        parsed_content = call_doubao_api(desensitized_content)
        print(f"[UPLOAD] OK 豆包 AI 返回，内容长度：{len(parsed_content)}")
        print(f"[UPLOAD] 步骤 5/5: AI 解析内容向量入库...")
        ai_chunk_count = add_to_chroma(parsed_content, file.filename, chunk_prefix="ai_")
        print(f"[UPLOAD] OK AI 解析内容入库，片段数：{ai_chunk_count}")
        file_record = create_file_record(
            filename=file.filename,
            file_size=len(file_content),
            file_type=file.filename.split(".")[-1].lower(),
            chunk_count=raw_chunk_count + ai_chunk_count,
            content_preview=parsed_content
        )
        file_path = save_uploaded_file(file_content, file.filename, file_record["id"])
        file_record["file_path"] = file_path
        add_file_record(file_record)
        print(f"[UPLOAD] OK 文件保存成功：{file_path}")
        print(f"\n{'='*60}")
        print(f"[UPLOAD] OK 上传完成：原始 {raw_chunk_count} 个片段 + AI {ai_chunk_count} 个片段 = 共 {raw_chunk_count + ai_chunk_count} 个")
        print(f"{'='*60}\n")
        return JSONResponse({
            "code": 200, "msg": "知识库搭建成功", "data": {
                "file_name": file.filename,
                "parsed_content": parsed_content,
                "chunk_count": raw_chunk_count + ai_chunk_count,
                "file_id": file_record["id"],
                "raw_chunks": raw_chunk_count,
                "ai_chunks": ai_chunk_count
            }
        })
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"整体流程失败：{str(e)}")


# ---------------------- 6. 知识库查询接口 ----------------------
@router.get("/query")
async def query_knowledge(query: str, top_k: int = 10):
    try:
        vector_model = _get_model()
        collection = _get_collection()
        query_embedding = vector_model.encode(query).tolist()
        results = collection.query(
            query_embeddings=[query_embedding],
            n_results=top_k,
            include=["documents", "metadatas", "distances"]
        )
        knowledge_items = []
        if results and results["documents"] and len(results["documents"]) > 0:
            for i, doc in enumerate(results["documents"][0]):
                metadata = results["metadatas"][0][i] if results["metadatas"] else {}
                distance = results["distances"][0][i] if results["distances"] else 0
                query_keywords = query.lower().split()
                doc_lower = doc.lower()
                keyword_matches = sum(1 for kw in query_keywords if kw in doc_lower)
                keyword_score = keyword_matches / len(query_keywords) if query_keywords else 0
                if 'sql' in query.lower() and metadata.get('has_sql'):
                    keyword_score += 0.3
                if metadata.get('source_type') == 'raw':
                    keyword_score += 0.2
                similarity_score = 1 - (distance / 2)
                final_score = 0.6 * similarity_score + 0.4 * keyword_score
                knowledge_items.append({
                    "content": doc,
                    "metadata": metadata,
                    "distance": distance,
                    "keyword_score": keyword_score,
                    "final_score": final_score
                })
        knowledge_items.sort(key=lambda x: x["final_score"], reverse=True)
        top_results = knowledge_items[:5]
        return JSONResponse({"code": 200, "msg": "查询成功", "data": {"query": query, "results": top_results}})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"查询失败：{str(e)}")


# ---------------------- 7. 知识库统计信息 ----------------------
@router.get("/stats")
async def get_knowledge_stats():
    try:
        collection = _get_collection()
        count = collection.count()
        return JSONResponse({"code": 200, "msg": "获取成功", "data": {"total_chunks": count, "collection_name": "dba_knowledge"}})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"获取统计信息失败：{str(e)}")


# ---------------------- 8. 文件管理接口 ----------------------
@router.get("/files")
async def list_files():
    try:
        files = get_all_files()
        return JSONResponse({"code": 200, "msg": "获取成功", "data": files})
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"获取文件列表失败：{str(e)}")


@router.get("/files/{file_id}")
async def get_file(file_id: str):
    try:
        file = get_file_record(file_id)
        if not file:
            raise HTTPException(status_code=404, detail="文件不存在")
        return JSONResponse({"code": 200, "msg": "获取成功", "data": file})
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"获取文件详情失败：{str(e)}")


@router.get("/files/{file_id}/download")
async def download_file(file_id: str):
    try:
        from pathlib import Path
        from urllib.parse import quote
        file_record = get_file_record(file_id)
        if not file_record:
            raise HTTPException(status_code=404, detail="文件不存在")
        file_path = Path(file_record["file_path"])
        if not file_path.exists():
            raise HTTPException(status_code=404, detail="文件已丢失")
        def iterfile():
            with open(file_path, mode="rb") as f:
                while chunk := f.read(1024 * 1024):
                    yield chunk
        filename = file_record["file_name"]
        encoded_filename = quote(filename)
        return StreamingResponse(
            iterfile(),
            media_type="application/octet-stream",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"下载文件失败：{str(e)}")


@router.delete("/files/{file_id}")
async def delete_file(file_id: str):
    try:
        file_record = get_file_record(file_id)
        if not file_record:
            raise HTTPException(status_code=404, detail="文件不存在")
        delete_physical_file(file_record)
        delete_file_record(file_id)
        return JSONResponse({"code": 200, "msg": "删除成功"})
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"删除文件失败：{str(e)}")
